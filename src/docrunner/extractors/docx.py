"""DOCX → Markdown by walking paragraphs, headings, and tables."""

from __future__ import annotations

import io

import docx as python_docx

from ..models import ExtractError, FetchedContent, ScrapeResult, SourceKind


def _para_to_md(para) -> str:
    text = para.text.strip()
    if not text:
        return ""
    style = (para.style.name or "").lower() if para.style else ""
    if style.startswith("heading"):
        # "Heading 2" -> level 2
        digits = "".join(ch for ch in style if ch.isdigit())
        level = int(digits) if digits else 1
        level = min(max(level, 1), 6)
        return f"{'#' * level} {text}"
    if style.startswith("list") or style.startswith("bullet"):
        return f"- {text}"
    return text


def _table_to_md(table) -> str:
    rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells]
            for row in table.rows]
    if not rows:
        return ""
    header, *body = rows
    width = len(header)
    lines = ["| " + " | ".join(header) + " |",
             "| " + " | ".join(["---"] * width) + " |"]
    for r in body:
        cells = (r + [""] * width)[:width]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def to_markdown(content: FetchedContent) -> ScrapeResult:
    try:
        document = python_docx.Document(io.BytesIO(content.data))
    except Exception as exc:
        raise ExtractError(f"Could not open DOCX: {exc}") from exc

    blocks: list[str] = []
    for para in document.paragraphs:
        chunk = _para_to_md(para)
        if chunk:
            blocks.append(chunk)
    for table in document.tables:
        chunk = _table_to_md(table)
        if chunk:
            blocks.append(chunk)

    title = ""
    cp = document.core_properties
    if cp and cp.title:
        title = cp.title.strip()
    if not title:
        # first heading or first line
        for b in blocks:
            if b.startswith("#"):
                title = b.lstrip("# ").strip()
                break
        if not title and blocks:
            title = blocks[0][:120]

    markdown = "\n\n".join(blocks).strip()
    warnings = [] if markdown else ["no extractable text in document"]
    return ScrapeResult(
        markdown=markdown,
        title=title,
        source_url=content.url,
        source_kind=SourceKind.DOCX,
        warnings=warnings,
    )
