import io

from docrunner.extractors import html as html_ex
from docrunner.extractors import docx as docx_ex
from docrunner.extractors import text as text_ex
from docrunner.models import FetchedContent, SourceKind


def _content(data: bytes, ct: str, url="https://example.com/x", name=""):
    return FetchedContent(url=url, final_url=url, content_type=ct, data=data, filename=name)


def test_html_extraction_title_and_body():
    html = b"""<html><head><title>My Page</title></head>
    <body><article><h1>Hello</h1><p>This is the body text of the article.</p>
    <a href="/doc.pdf">a pdf</a></article></body></html>"""
    res = html_ex.to_markdown(_content(html, "text/html"))
    assert res.title == "My Page"
    assert "body text of the article" in res.markdown
    assert any(link.endswith("/doc.pdf") for link in res.links)
    assert res.source_kind == SourceKind.WEBPAGE


def test_text_passthrough_collapses_blank_lines():
    raw = b"# Title\n\n\n\nsome content\r\n"
    res = text_ex.to_markdown(_content(raw, "text/plain"))
    assert "some content" in res.markdown
    assert "\n\n\n" not in res.markdown
    assert res.title == "Title"


def test_docx_extraction_roundtrip():
    import docx as python_docx

    doc = python_docx.Document()
    doc.add_heading("Report Title", level=1)
    doc.add_paragraph("First paragraph.")
    doc.add_heading("Section", level=2)
    doc.add_paragraph("Second paragraph.")
    buf = io.BytesIO()
    doc.save(buf)

    res = docx_ex.to_markdown(
        _content(
            buf.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    )
    assert "# Report Title" in res.markdown
    assert "## Section" in res.markdown
    assert "First paragraph." in res.markdown
    assert res.source_kind == SourceKind.DOCX


def test_pdf_extraction():
    # Build a one-page PDF with text using pypdf's writer + reportlab-free approach.
    # pypdf can't author text content, so use a minimal hand-built PDF.
    pdf_bytes = _minimal_pdf("Hello PDF World")
    from docrunner.extractors import pdf as pdf_ex

    res = pdf_ex.to_markdown(_content(pdf_bytes, "application/pdf", name="doc.pdf"))
    # Either we extract the text, or we warn cleanly — never crash.
    assert res.source_kind == SourceKind.PDF
    assert isinstance(res.markdown, str)


def _minimal_pdf(text: str) -> bytes:
    """A tiny valid single-page PDF containing one text string."""
    objs = []
    objs.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objs.append(b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    objs.append(
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
    )
    stream = b"BT /F1 24 Tf 72 700 Td (" + text.encode() + b") Tj ET"
    objs.append(b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream))
    objs.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, body in enumerate(objs, start=1):
        offsets.append(len(out))
        out += b"%d 0 obj\n" % i + body + b"\nendobj\n"
    xref_pos = len(out)
    out += b"xref\n0 %d\n" % (len(objs) + 1)
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += b"%010d 00000 n \n" % off
    out += b"trailer\n<< /Size %d /Root 1 0 R >>\n" % (len(objs) + 1)
    out += b"startxref\n%d\n%%%%EOF" % xref_pos
    return bytes(out)
